import re
import os
import anthropic
from youtube_transcript_api import YouTubeTranscriptApi
from anthropic import Anthropic, HUMAN_PROMPT, AI_PROMPT
import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from tkinter import ttk
from docx import Document
import threading
from ttkthemes import ThemedTk
from PIL import Image, ImageTk
import time
import random

# Set up Claude API key
claude_api_key = "claudeai api key"

# Instantiate the Claude client
client = Anthropic(api_key=claude_api_key)

# Function to extract video ID
def extract_video_id(youtube_url):
    video_id_regex = r'(?:v=|\/)([0-9A-Za-z_-]{11}).*'
    match = re.search(video_id_regex, youtube_url)
    if match:
        return match.group(1)
    return None

# Function to fetch transcript
def fetch_transcript(video_id):
    try:
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        text_list = [item['text'] for item in transcript]
        transcript_text = ' '.join(text_list)
        return transcript_text
    except Exception as e:
        messagebox.showerror("Error", f"Error fetching transcript: {e}")
        return None

# Function to summarize transcript using Claude API
def summarize_claude(text):
    try:
        system_prompt = "You are a helpful assistant that summarizes YouTube videos. Summarize the following lecture in a concise manner."
        prompt = f"{HUMAN_PROMPT} {system_prompt} {text}{AI_PROMPT}"
        response = client.completions.create(
            model="claude-2",
            prompt=prompt,
            max_tokens_to_sample=300
        )
        summary = response.completion.strip()
        return summary
    except Exception as e:
        messagebox.showerror("Error", f"Error summarizing text: {e}")
        return None

# Function to translate text using Claude API
def translate_claude(text, target_language):
    try:
        system_prompt = f"You are a helpful assistant that translates text. Translate the following text to {target_language}."
        prompt = f"{HUMAN_PROMPT} {system_prompt} {text}{AI_PROMPT}"
        response = client.completions.create(
            model="claude-2",
            prompt=prompt,
            max_tokens_to_sample=300
        )
        translation = response.completion.strip()
        return translation
    except Exception as e:
        messagebox.showerror("Error", f"Error translating text: {e}")
        return None

# Function to provide detailed explanation with figures and examples and save it to Word file
def explain_with_figures(text):
    try:
        system_prompt = "You are a helpful assistant that provides detailed explanations of the following topic with figures and examples."
        prompt = f"{HUMAN_PROMPT} {system_prompt} {text}{AI_PROMPT}"
        response = client.completions.create(
            model="claude-2",
            prompt=prompt,
            max_tokens_to_sample=500
        )
        explanation = response.completion.strip()
        save_explanation_to_word(explanation)
        return explanation
    except Exception as e:
        messagebox.showerror("Error", f"Error providing explanation: {e}")
        return None

# Function to save explanation to Word file
def save_explanation_to_word(explanation):
    try:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], initialfile="explanation")
        if file_path:
            document = Document()
            document.add_heading('Detailed Explanation', level=1)
            document.add_paragraph(explanation)
            document.save(file_path)
            messagebox.showinfo("Success", "Explanation saved successfully!")
    except Exception as e:
        messagebox.showerror("Error", f"Error saving explanation: {e}")

# Function to save summary to Word file
def save_summary_to_word(summary):
    try:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], initialfile="summary")
        if file_path:
            document = Document()
            document.add_heading('YouTube Video Summary', level=1)
            document.add_paragraph(summary)
            document.save(file_path)
            messagebox.showinfo("Success", "Summary saved successfully!")
    except Exception as e:
        messagebox.showerror("Error", f"Error saving summary: {e}")

# Function to save translated summary to Word file
def save_translated_summary_to_word(summary, language):
    try:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], initialfile=f"summary_{language}")
        if file_path:
            document = Document()
            document.add_heading(f'YouTube Video Summary ({language})', level=1)
            document.add_paragraph(summary)
            document.save(file_path)
            messagebox.showinfo("Success", f"Summary saved successfully in {language}!")
    except Exception as e:
        messagebox.showerror("Error", f"Error saving summary: {e}")

# Main function to create UI
def main():
    def on_summarize():
        youtube_url = url_entry.get()
        video_id = extract_video_id(youtube_url)

        if video_id:
            loading_label.pack(pady=15)
            spinner.pack(pady=10)
            threading.Thread(target=fetch_and_summarize, args=(video_id,)).start()
        else:
            messagebox.showerror("Error", "Could not extract video ID. Please check the URL and try again.")

    def fetch_and_summarize(video_id):
        try:
            for i in range(3):
                loading_label.config(text=f"Summarizing, please wait{'.' * (i % 3 + 1)}")
                loading_label.update_idletasks()
                time.sleep(0.5)
            transcript_text = fetch_transcript(video_id)
            if transcript_text:
                summary = summarize_claude(transcript_text)
                if summary:
                    summary_text.delete(1.0, tk.END)
                    summary_text.insert(tk.END, summary)
                    save_button.config(state=tk.NORMAL)
                    translate_button.config(state=tk.NORMAL)
                    explain_button.config(state=tk.NORMAL)
                else:
                    summary_text.delete(1.0, tk.END)
                    save_button.config(state=tk.DISABLED)
                    translate_button.config(state=tk.DISABLED)
                    explain_button.config(state=tk.DISABLED)
            else:
                summary_text.delete(1.0, tk.END)
                save_button.config(state=tk.DISABLED)
                translate_button.config(state=tk.DISABLED)
                explain_button.config(state=tk.DISABLED)
        finally:
            loading_label.pack_forget()
            spinner.pack_forget()

    def on_translate():
        target_language = language_var.get()
        summary = summary_text.get(1.0, tk.END).strip()
        if summary:
            translated_summary = translate_claude(summary, target_language)
            if translated_summary:
                save_translated_summary_to_word(translated_summary, target_language)

    def on_explain():
        text = summary_text.get(1.0, tk.END).strip()
        if text:
            explanation = explain_with_figures(text)
            if explanation:
                summary_text.delete(1.0, tk.END)
                summary_text.insert(tk.END, explanation)

    def toggle_dark_mode():
        nonlocal is_dark_mode
        is_dark_mode = not is_dark_mode
        if is_dark_mode:
            root.configure(bg="#181818")
            canvas.configure(bg="#181818")
            title_label.configure(fg="#FFFFFF", bg="#181818")
            url_label.configure(fg="#FFCC00", bg="#181818")
            summary_text.configure(bg="#282828", fg="#E0E0E0")
            loading_label.configure(fg="#FFCC00", bg="#181818")
            dark_mode_button.configure(text="Light Mode")
        else:
            root.configure(bg="#FFFFFF")
            canvas.configure(bg="#FFFFFF")
            title_label.configure(fg="#000000", bg="#FFFFFF")
            url_label.configure(fg="#00008B", bg="#FFFFFF")
            summary_text.configure(bg="#F5F5F5", fg="#000000")
            loading_label.configure(fg="#00008B", bg="#FFFFFF")
            dark_mode_button.configure(text="Dark Mode")

    # Set up the main application window
    root = ThemedTk(theme="breeze")
    root.title("AI YouTube Summarizer")
    root.configure(bg="#FFFFFF")
    root.geometry("800x600")
# Add this code for full-screen functionality
    root.attributes('-fullscreen', True)
    root.bind("<Escape>", lambda event: root.attributes('-fullscreen', False))
    is_dark_mode = False

    # Create a background animation using canvas
    canvas = tk.Canvas(root, width=800, height=600, bg="#FFFFFF", highlightthickness=0)
    canvas.place(x=0, y=0, relwidth=1, relheight=1)

    # Create animated circles for the background
    circles = []
    for _ in range(15):
        x = random.randint(0, 800)
        y = random.randint(0, 600)
        size = random.randint(20, 60)
        circle = canvas.create_oval(x, y, x + size, y + size, outline="", fill="#ADD8E6", stipple="gray12")
        circles.append((circle, x, y))

    def animate_circles():
        for circle, x, y in circles:
            new_x = x + random.randint(-2, 2)
            new_y = y + random.randint(-2, 2)
            canvas.move(circle, new_x - x, new_y - y)
        root.after(100, animate_circles)

    animate_circles()

    # Title Label with Animation
    title_text = "AI YouTube Summarizer"
    title_label = tk.Label(root, text="", fg="#000000", bg="#FFFFFF", font=("Arial", 24, "bold"))
    title_label.pack(pady=10)

    def animate_title():
        for i in range(len(title_text) + 1):
            title_label.config(text=title_text[:i])
            root.update()
            time.sleep(0.1)

    threading.Thread(target=animate_title).start()

    # URL Entry
    url_label = tk.Label(root, text="Enter YouTube Video URL:", fg="#00008B", bg="#FFFFFF", font=("Arial", 14, "bold"))
    url_label.pack(pady=10)
    url_entry = tk.Entry(root, width=50, font=("Arial", 14), relief="flat", borderwidth=2)
    url_entry.pack(pady=5)
    url_entry.config(highlightbackground="#00008B", highlightcolor="#00008B", highlightthickness=1)

    # Summarize Button with Animation
    def button_hover_in(event):
        summarize_button.config(bg="#FF3333")

    def button_hover_out(event):
        summarize_button.config(bg="#FF0000")

    summarize_button = tk.Button(root, text="Summarize", command=on_summarize, bg="#FF0000", fg="white", font=("Arial", 14, "bold"), activebackground="#CC0000")
    summarize_button.pack(pady=10)
    summarize_button.bind("<Enter>", button_hover_in)
    summarize_button.bind("<Leave>", button_hover_out)

    # Loading Spinner
    spinner = ttk.Progressbar(root, mode='indeterminate', length=200)
    loading_label = tk.Label(root, text="Summarizing, please wait...", fg="#00008B", bg="#FFFFFF", font=("Arial", 12))

    # Summary Text Box
    summary_text = tk.Text(root, height=10, width=70, wrap=tk.WORD, bg="#F5F5F5", fg="#000000", font=("Arial", 14), relief="flat")
    summary_text.pack(pady=10)
    summary_text.insert(tk.END, "Your summary will appear here...")
    summary_text.bind("<FocusIn>", lambda e: summary_text.delete(1.0, tk.END))

    # Save Button with Animation
    def save_button_hover_in(event):
        save_button.config(bg="#33A1FF")

    def save_button_hover_out(event):
        save_button.config(bg="#008CBA")

    save_button = tk.Button(root, text="Save Summary to Word", command=lambda: save_summary_to_word(summary_text.get(1.0, tk.END)), bg="#008CBA", fg="#FFFFFF", font=("Arial", 14, "bold"), state=tk.DISABLED, activebackground="#007bb5")
    save_button.pack(pady=10)
    save_button.bind("<Enter>", save_button_hover_in)
    save_button.bind("<Leave>", save_button_hover_out)

    # Language Selection Dropdown
    language_var = tk.StringVar(root)
    language_var.set("English")
    languages = ["English", "Russian", "Urdu"]
    language_menu = ttk.Combobox(root, textvariable=language_var, values=languages, state="readonly", font=("Arial", 14))
    language_menu.pack(pady=10)

    # Translate Button
    translate_button = tk.Button(root, text="Translate and Save", command=on_translate, bg="#FF6347", fg="white", font=("Arial", 14, "bold"), state=tk.DISABLED, activebackground="#FF4500")
    translate_button.pack(pady=10)

    # Explain Button (moved to the right of the screen)
    explain_button = tk.Button(root, text="Explain Topic", command=on_explain, bg="#32CD32", fg="white", font=("Arial", 14, "bold"), state=tk.DISABLED, activebackground="#2E8B57")
    explain_button.place(relx=0.85, rely=0.5, anchor=tk.CENTER)

    # Dark Mode Toggle Button
    dark_mode_button = tk.Button(root, text="Dark Mode", command=toggle_dark_mode, bg="#808080", fg="white", font=("Arial", 12, "bold"), activebackground="#696969")
    dark_mode_button.place(relx=0.05, rely=0.05, anchor=tk.NW)

    root.mainloop()

if __name__ == "__main__":
    main()